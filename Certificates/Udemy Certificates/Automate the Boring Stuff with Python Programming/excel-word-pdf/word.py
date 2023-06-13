import docx;

path = 'excel-word-pdf/docs/exercise.docx';
doc = docx.Document(path);
#doc.paragraphs
for i in range(len(doc.paragraphs)):
    print(doc.paragraphs[i].text);

paragraph = doc.paragraphs[0];
paragraph.text = "python invaders";
paragraph.add_run(" are here");
paragraph.runs[1].bold = True;
paragraph.underline = True;
doc.add_paragraph("why am i here");
doc.save("new.docx");

def getText(file):
    doc = docx.Document(file);
    text = [];
    
    for paragraph in doc.paragraphs:
        text.append(paragraph.text);
    return "\n".join(text);

print(getText(path));